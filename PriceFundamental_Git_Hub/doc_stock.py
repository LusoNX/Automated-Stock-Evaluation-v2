from __future__ import print_function
from mailmerge import MailMerge
from datetime import datetime
import numpy as np
from dateutil.relativedelta import relativedelta
import time
import requests
import sqlite3
from selenium import webdriver  
from bs4 import BeautifulSoup
import random
from scipy import stats
from scipy.stats import norm
from pandas.io.json import json_normalize
import requests
import urllib
from sqlalchemy import create_engine
import pyodbc
import pandas as pd
from selenium import webdriver
from bs4 import BeautifulSoup
from selenium.webdriver.support.ui import Select
from selenium.webdriver.common.by import By
from selenium.webdriver.common.action_chains import ActionChains
import time
import random
from datetime import date
import investpy
import yfinance as yf
from datetime import date
from dateutil.relativedelta import relativedelta
import seaborn as sns
import matplotlib.pyplot as plt
from scipy import stats
from scipy.stats import norm
import statsmodels.api as sm
from sklearn.preprocessing import PolynomialFeatures
import operator
import sys
import stock_evaluation
from docx import Document
import fundamental_evaluation

import os
template = "company-presentation-template.docx"

document = MailMerge(template)

conn_str = (r'DRIVER={Microsoft Access Driver (*.mdb, *.accdb)};'
			r'DBQ=YOUR DIRECTORY\StockPriceFundamentalData.accdb;')
cnn_url = f"access+pyodbc:///?odbc_connect={urllib.parse.quote_plus(conn_str)}"
acc_engine = create_engine(cnn_url)

def create_directory(symbol):
	folder_name = symbol
	existing_directory_2 = os.path.isdir(r"YOUR DIRECTORY\{}".format(symbol))
	if existing_directory_2 == True:
		pass
	else:
		folder = os.path.join(r'YOUR DIRECTORY', symbol)
		os.makedirs(folder)
	folder = os.path.join(r'\{}'.format(symbol), folder_name)
	os.makedirs(folder)
 




def add_picture(directory,_doc,graph_name,file_name,stock_symbol):
	graph_name = "[{}]".format(graph_name)
	graph = [i for i, p in enumerate(_doc.paragraphs) if graph_name in p.text] ## Use this to identify the paragraph of a specific wording
	p = _doc.paragraphs[graph[0]]
	p.text = ""
	r = p.add_run()
	try:
		r.add_picture(directory+r"\{}_{}.png".format(file_name,stock_symbol))
	except FileNotFoundError:
		print("File Not FOund for graph 1")
		pass

def document_reader(id_stock):
	df_stock_index = pd.read_sql("SELECT * FROM StockIndex WHERE ID_STOCK = {}".format(id_stock),acc_engine)
	id_industry = pd.read_sql("SELECT (ID_INDUSTRY) FROM IndustryStockMerge WHERE ID_STOCK = {}".format(id_stock),acc_engine)["ID_INDUSTRY"].values[0]
	df_ratio_index = pd.read_sql("SELECT * FROM RatioData WHERE ID_STOCK = {}".format(id_stock),acc_engine).sort_values("Data_Appended").iloc[[-1]]
	df_price = pd.read_sql("SELECT Data,Close FROM PriceData WHERE ID_STOCK = {}".format(id_stock),acc_engine).sort_values("Data")
	last_price = df_price["Close"].iloc[-1]
	stock_name = df_stock_index["STOCK_NAME"].iloc[0]
	stock_symbol =df_stock_index["SYMBOL"].iloc[0]
	_industry = str(pd.read_sql("SELECT (NAME_INDUSTRY) FROM IndustryIndex WHERE ID_INDUSTRY = {}".format(id_industry),acc_engine)["NAME_INDUSTRY"].values[0])
	df_ratio_index = df_ratio_index.round(2)
	df_ratio_index = df_ratio_index.astype(str)

	target_recommendation = df_ratio_index["target_mean"].iloc[0]
	_nr_analyst =df_ratio_index["number_of_analyst"].iloc[0]
	_esg_score = df_ratio_index["esg_score"].iloc[0]
	div_yield = str(round(float(df_ratio_index["div_rate"].iloc[0]),2)*100)+"%"
	_short_ratio = df_ratio_index["short_ratio"].iloc[0]
	#Append the introductory data
	document.merge(ticker = stock_symbol)
	document.merge(name = stock_name)
	#document.merge(description = _description)
	document.merge(industry = _industry)
	document.merge(analyst_consensus = target_recommendation)
	document.merge(nr_analyst = _nr_analyst)
	document.merge(esg_score = str(_esg_score))
	document.merge(short_ratio = str(_short_ratio))

	existing_directory = os.path.isdir(r"YOUR DIRECTORY\{}".format(stock_symbol))
	if existing_directory == True:
		pass
	else:
		create_directory(stock_symbol)

	#document.merge(country = _country)
	gross_margin = df_ratio_index["gross_margin"].iloc[0]
	operating_margin = df_ratio_index["operating_margin"].iloc[0]
	ni_margin = df_ratio_index["ni_margin"].iloc[0]
	_roa = df_ratio_index["roa"].iloc[0]
	_roe = df_ratio_index["roe"].iloc[0]
	pe_ratio = df_ratio_index["pe_ratio"].iloc[0]
	f_pe_ratio =df_ratio_index["forward_pe_ratio"].iloc[0]
	peg_ratio = df_ratio_index["peg_ratio"].iloc[0]
	ev_to_ebitda =df_ratio_index["ev_to_ebitda"].iloc[0]
	pb_ratio =df_ratio_index["price_to_book_ratio"].iloc[0]
	price_to_cfo =df_ratio_index["price_to_cfo"].iloc[0]
	price_to_fcf =df_ratio_index["price_to_fcf"].iloc[0]
	_price_to_sales = df_ratio_index["price_to_sales"].iloc[0]
	debt_to_equity =df_ratio_index["debt_to_equity"].iloc[0]
	cash_to_assets =df_ratio_index["cash_to_assets"].iloc[0]
	_coverage_ratio = df_ratio_index["coverage_ratio"].iloc[0]
	_coverage_ratio_prime =df_ratio_index["coverage_ratio_prime"].iloc[0]
	_div_rate =df_ratio_index["div_rate"].iloc[0]

	# COmpany ratios DOC
	document.merge(gross_company = gross_margin)
	document.merge(operating_company = operating_margin)
	document.merge(ni_company = ni_margin)
	document.merge(roa_company = _roa)
	document.merge(roe_company = _roe)
	document.merge(de_company = debt_to_equity)
	document.merge(pe_company = pe_ratio)
	document.merge(pb_company = pb_ratio)
	document.merge(peg_ratio = peg_ratio)
	document.merge(forward_pe = f_pe_ratio)
	document.merge(_price_cfo = price_to_cfo)
	document.merge(price_fcf = price_to_fcf)
	document.merge(price_to_sales = _price_to_sales)
	document.merge(price_fcf = price_to_fcf)
	document.merge(cash_assets =cash_to_assets)
	document.merge(ev_to_ebitda_company = ev_to_ebitda)
	document.merge(coverage_ratio = _coverage_ratio)
	document.merge(coverage_prime = _coverage_ratio_prime)
	document.merge(div_rate = _div_rate)

	df_industry_ratio = pd.read_sql("SELECT * FROM IndustryAverage WHERE ID_INDUSTRY = {}".format(id_industry),acc_engine).sort_values("Data_Appended")
	df_industry_ratio = df_industry_ratio.iloc[[-1]]
	df_industry_ratio = df_industry_ratio.round(2)
	df_industry_ratio = df_industry_ratio.astype(str)

	# Industry Ratios DOC
	document.merge(gross_industry = df_industry_ratio["gross_margin_peer"].values[0])
	document.merge(operating_industry = df_industry_ratio["operating_margin_peer"].values[0])
	document.merge(ni_industry =  df_industry_ratio["ni_margin_peer"].values[0])
	document.merge(roa_industry =  df_industry_ratio["roa_peer"].values[0])
	document.merge(roe_industry =  df_industry_ratio["roe_peer"].values[0])
	document.merge(de_industry =  df_industry_ratio["debt_to_equity_peer"].values[0])
	document.merge(coverage_ratio_industry =  df_industry_ratio["debt_to_equity_peer"].values[0])
	document.merge(coverage_ratio_prime_industry =  df_industry_ratio["debt_to_equity_peer"].values[0])
	document.merge(div_rate_industry =  df_industry_ratio["div_rate_peer"].values[0])

	document.merge(pe_industry =  df_industry_ratio["pe_ratio_peer"].values[0])
	document.merge(pb_industry =  df_industry_ratio["price_to_book_peer"].values[0])
	document.merge(peg_industry =  df_industry_ratio["peg_ratio_peer"].values[0])
	document.merge(p_to_sales_industry =  df_industry_ratio["price_to_sales_peer"].values[0])
	document.merge(ev_to_ebitda_industry =  df_industry_ratio["ev_to_ebitda_peer"].values[0])
	document.merge(cash_assets_industry =  df_industry_ratio["cash_to_assets_peer"].values[0])
	document.merge(p_to_cfo_industry =  df_industry_ratio["price_to_cfo_peer"].values[0])
	document.merge(p_to_fcf_industry =  df_industry_ratio["price_to_fcf_peer"].values[0])
	document.merge(f_pe_industry =  df_industry_ratio["forward_pe_ratio_peer"].values[0])



	#document.merge(cost_of_equity = ke)
	#document.merge(wacc = wacc)

	list_values_market =stock_evaluation.get_beta_figures(id_stock,"Market 3 Moment",156)
	list_values_carhart = stock_evaluation.get_beta_figures(id_stock,"Carhart",156)

	document.merge(last_beta_market = list_values_market[2])
	document.merge(mean_beta_market = list_values_market[0])
	document.merge(std_beta_market = list_values_market[1])
	document.merge(beta_cosk_market_2 = list_values_market[3])
	document.merge(last_r2 = list_values_market[6])
	document.merge(last_total_beta = list_values_market[4])
	document.merge(mkt_beta = list_values_carhart[0])
	document.merge(sml_carhart = list_values_carhart[1])
	document.merge(hml_carhart = list_values_carhart[2])
	document.merge(momentum_carhart = list_values_carhart[3])
	document.merge(r2_carhart = list_values_carhart[4])
	document.merge(sml_pvalues = list_values_carhart[5])
	document.merge(hml_pvalues = list_values_carhart[6])
	document.merge(momentum_pvalues = list_values_carhart[7])



	# Industry average valuation rations
	ddm_1= stock_evaluation.cash_flow_valuation_model(id_stock,"Market 3 Moment","DDM",1000)
	ddm_2= stock_evaluation.cash_flow_valuation_model(id_stock,"Carhart","DDM",1000)
	ddm_3= stock_evaluation.cash_flow_valuation_model(id_stock,"Industry Fama and French","DDM",1000)
	fcf_1= stock_evaluation.cash_flow_valuation_model(id_stock,"Market 3 Moment","FCF",1000)
	fcf_2= stock_evaluation.cash_flow_valuation_model(id_stock,"Carhart","FCF",1000)
	fcf_3= stock_evaluation.cash_flow_valuation_model(id_stock,"Industry Fama and French","FCF",1000)
	fcfe_1= stock_evaluation.cash_flow_valuation_model(id_stock,"Market 3 Moment","FCFE",1000)
	fcfe_2= stock_evaluation.cash_flow_valuation_model(id_stock,"Carhart","FCFE",1000)
	fcfe_3= stock_evaluation.cash_flow_valuation_model(id_stock,"Industry Fama and French","FCFE",1000)

	
	if ddm_1 == 0:
		mean_1 = np.mean([fcf_1,fcf_2,fcf_3])
		mean_2 = np.mean([fcfe_1,fcfe_2,fcfe_3])
		mean_3 = np.mean([ddm_1,ddm_2,ddm_3])
		final_mean_val = np.mean([mean_1,mean_2])

	else:
		mean_1 = np.mean([fcf_1,fcf_2,fcf_3])
		mean_2 = np.mean([fcfe_1,fcfe_2,fcfe_3])
		mean_3 = np.mean([ddm_1,ddm_2,ddm_3])
		final_mean_val = np.mean([mean_1,mean_2,mean_3])



	fcf_1 = str(round(fcf_1,2))
	fcf_2 = str(round(fcf_2,2))
	fcf_3 = str(round(fcf_3,2))
	fcfe_1 = str(round(fcfe_1,2))
	fcfe_2 = str(round(fcfe_2,2))
	fcfe_3 = str(round(fcfe_3,2))
	ddm_1 = str(round(ddm_1,2))
	ddm_2 = str(round(ddm_2,2))
	ddm_3 =str(round(ddm_3,2))
	mean_1 =str(round(mean_1,2))
	mean_2 =str(round(mean_2,2))
	mean_3=str(round(mean_3,2))
	final_mean_val =str(round(final_mean_val,2))


	document.merge(dcf_fcf_3_moment = fcf_1)
	document.merge(dcf_fcf_carhart = fcf_2)
	document.merge(dcf_fcf_industry = fcf_3)
	document.merge(dcf_fcfe_3_moment = fcfe_1)
	document.merge(dcf_fcfe_carhart = fcfe_2)
	document.merge(dcf_fcfe_industry = fcfe_3)
	document.merge(ddm_3_moment =ddm_1)
	document.merge(ddm_carhart = ddm_2)
	document.merge(ddm_industry = ddm_3)
	document.merge(mean_dcf_fcf = mean_1)
	document.merge(mean_dcf_fcfe = mean_2)
	document.merge(mean_ddm = mean_3)
	document.merge(dcf_final_mean = final_mean_val)

	# Industry average valuation rations
	list_values_ind =  stock_evaluation.peer_valuation(id_stock)
	document.merge(value_ps = str(list_values_ind[0]))
	document.merge(value_pb = str(list_values_ind[1]))
	document.merge(value_pe = str(list_values_ind[2]))
	document.merge(value_cfo = str(list_values_ind[3]))
	document.merge(value_fcf = str(list_values_ind[4]))
	document.merge(mean_peer = str(list_values_ind[5]))
	document.merge(ind_g_peg = str(list_values_ind[6]))

	mean_target = round((float(final_mean_val)+list_values_ind[6]),2)
	if mean_target*1.05 > last_price: # Band created to differentiate between a buy and a hold. For example. a stock valued at 25*(1.10) 27.5 with a price of 28, is worth the buying because the upside return is above at least 10 %, while at a price of 26 the dollar difference is more likea hold
		target_recommendation = "BUY"
	elif mean_target > last_price:
		target_recommendation = "HOLD"
	elif mean_target*1.05 < last_price:# Same logic but the opos.
		target_recommendation = "SELL"
	else:
		target_recommendation = "HOLD"

	upside_downside_var = round(((mean_target - last_price)/last_price*100),2)

	

	document.merge(up_down_var = str(upside_downside_var) + "%")
	document.merge(target_price =str(mean_target))
	document.merge(recommendation = target_recommendation)
	# get regression qr 

	df_qr = stock_evaluation.qr_regression_market(id_stock)
	document.write('test-output.docx')

	## THis second part is for the adding of immages to the document

	doc = Document("test-output.docx")
	directory = r"YOUR DIRECTORY\{}".format(stock_symbol)

	## 
	fundamental_evaluation.get_fundamental_ratios(id_stock)
	add_picture(directory,doc,"beta_r2_graph_1","beta_3_moment",stock_symbol)
	add_picture(directory,doc,"p_values_graph_1","p_values_3_moment",stock_symbol)
	add_picture(directory,doc,"beta_4_factor","beta_4_moment",stock_symbol)
	add_picture(directory,doc,"p_values_graph_2","pvalues_4_moment",stock_symbol)
	add_picture(directory,doc,"gross_ope_ni_margin","goi_margin",stock_symbol)
	add_picture(directory,doc,"cfo_margin","cfo_margin",stock_symbol)
	add_picture(directory,doc,"roa_roe","roa_roe",stock_symbol)
	add_picture(directory,doc,"coverage_ratio","coverage_ratio",stock_symbol)
	add_picture(directory,doc,"qr_regression_graph","qr_regression_market",stock_symbol)
	doc.save(directory +r"\{i} Summary.docx".format(i = stock_symbol))

document_reader(1)
